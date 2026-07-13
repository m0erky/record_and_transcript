"""Export von Transkripten als Word-Dokument."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document


def export_to_docx(
    text: str,
    output_path: Path,
    title: str | None = None,
    metadata: dict[str, str] | None = None,
) -> Path:
    document = Document()
    document.add_heading(title or "Audio-Transkript", level=1)

    if metadata:
        for key, value in metadata.items():
            document.add_paragraph(f"{key}: {value}")
        document.add_paragraph("")

    for paragraph in text.split("\n"):
        document.add_paragraph(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def default_title(timestamp: datetime | None = None) -> str:
    moment = timestamp or datetime.now()
    return f"Transkript {moment.strftime('%d.%m.%Y %H:%M')}"
